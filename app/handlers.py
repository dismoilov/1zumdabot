from aiogram.filters import CommandStart
import app.keyboards as kb
from aiogram import Router, F
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
import re
from aiogram.utils.keyboard import ReplyKeyboardBuilder
from .dicts import regions
from app.db import *
from aiogram.types import ReplyKeyboardRemove, Message
from datetime import datetime
import os
from aiogram.types import FSInputFile

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL


def generate_resume(data):
    doc = Document()

    # Устанавливаем стиль основного шрифта
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Таблица для фото и основной информации
    if data.get('face_photo'):
        photo_table = doc.add_table(rows=1, cols=2)
        photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells = photo_table.rows[0].cells
        row_cells[0].width = Inches(1.5)

        # Фото слева
        row_cells[0].paragraphs[0].add_run().add_picture(data['face_photo'], width=Inches(1.5))

        # Контактная информация справа
        info_cell = row_cells[1].paragraphs[0]
        info_cell.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        info_cell.add_run(f"{data['first_name']} {data['middle_name']} {data['last_name']}\n").bold = True
        info_cell.add_run(f"Telefon raqam: {data['phone']}\n")
        info_cell.add_run(f"Qo'sh. Telefon raqam: {data.get('extra_phone', '-')}\n")
        info_cell.add_run(f"Yashash joyi: {data.get('region')}, {data.get('district')}, {data.get('address')}\n")
        doc.add_paragraph()

    # Создаем таблицу для основных разделов информации
    sections_table = doc.add_table(rows=1, cols=3)
    sections_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Shaxsiy Ma'lumotlar
    personal_info_cell = sections_table.cell(0, 0)
    personal_info_cell.paragraphs[0].add_run("Shaxsiy Ma'lumotlar").bold = True
    personal_info_cell.add_paragraph(f"Jinsi: {data['gender']}")
    personal_info_cell.add_paragraph(f"Tug'ilgan sana: {data['birthday']}")
    personal_info_cell.add_paragraph(f"Oilaviy holati: {data['family_status']}")
    personal_info_cell.add_paragraph(f"O'zbekiston fuqarosi: {data['is_uzbek_citizen']}")

    # O'quv Ma'lumot
    education_info_cell = sections_table.cell(0, 1)
    education_info_cell.paragraphs[0].add_run("O'quv Ma'lumot").bold = True
    education_info_cell.add_paragraph(f"O'qish joyi: {data['education']} ({data['education_place']})")
    education_info_cell.add_paragraph(f"Hozirda talabami: {data['is_studying']} ({data['education_type']})")
    education_info_cell.add_paragraph(f"Mutaxassislik: {data['specialty']}")

    # Ish Tajribasi
    experience_info_cell = sections_table.cell(0, 2)
    experience_info_cell.paragraphs[0].add_run("Ish Tajribasi").bold = True
    experience_info_cell.add_paragraph(f"Tajriba: {data['experience']}")
    experience_info_cell.add_paragraph(f"Hozir ishlayaptimi: {data['is_working_now']}")
    experience_info_cell.add_paragraph(f"Ilgari bu kompaniyada ishlaganmi: {data['worked_before']}")

    # Отдельная таблица для "Qo'shimcha Ma'lumot"
    doc.add_paragraph("\nQo'shimcha Ma'lumot", style='Heading 1')
    extra_info_table = doc.add_table(rows=0, cols=2)
    extra_info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Добавляем строки в таблицу для каждого пункта с отступом в первой колонке
    extra_fields = [
        ("Istalgan lavozim", data['post']),
        ("Istalgan ish filiali", data['branch_desire']),
        ("Istalgan maosh", f"{data['salary']} so'm"),
        ("Yaqin tanishi bormi", data['is_familiar_works_here']),
        ("Biladigan tillar", ', '.join(data['languages'])),
        ("Biladigan dasturlar", ', '.join(data['programms_experience'])),
        ("Ish tajriba haqida qo'shimcha ma'lumotlar", data['work_experience']),
        ("Qo'shimcha ma'lumotlar", data['about'])
    ]

    for field, value in extra_fields:
        row = extra_info_table.add_row().cells
        row[0].text = f"{field}: "  # Добавляем двоеточие для более четкого разделения
        row[0].paragraphs[0].paragraph_format.left_indent = Inches(0.5)  # Отступ для текста первой колонки
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        row[1].text = value
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Устанавливаем ширину для колонок
        row[0].width = Inches(2.2)
        row[1].width = Inches(4.0)

    # Добавляем дату создания в нижний колонтитул
    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    current_datetime = datetime.now().strftime("Sana: %d-%m-%Y   Vaqt: %H:%M")
    footer.add_run(current_datetime).font.size = Pt(10)

    # Сохранение документа
    output_path = f'djangobot/media/resumes/{data["first_name"]}_{data["last_name"]}.docx'
    doc.save(output_path)

    return output_path


router = Router()

answer = "Salom 👋🏻\n\nUshbu bot 1Zumda do'konlar tarmog'iga ishga kirish maqsadida tashqi nomzodlar anketa to'ldirishi uchun mo'ljallangan!\n\nBu yerda siz kompaniyada mavjud ish o'rinlariga ariza sifatida ma'lumotlaringizni qoldirishingiz mumkin.\n\nDiqqat! Bot faqat faol bo'lgan bo'sh ish o'rinlariga tashqi nomzodlar anketa qoldirishlari uchun mo'ljallangan. Agar siz hozirda 1Zumda kompaniyasining xodimi bo'lsangiz va bo'sh ish o'rniga o'tmoqchi bo'lsangiz, ushbu bot sizga bu borada yordam bera olmaydi. Siz so'rovingiz bilan o'zingiz faoliyat yuritib turgan filialdagi bo'lim raxbaringizga murojaat qilishingiz kerak."

button_languages = ["O'zbek tili", "Ingliz tili", "Rus tili"]
button_programms = ["Word", "Excel", "PowerPoint"]


async def send_message_to_user(user_id: int, photo: str, caption: str, bot):
    try:
        await bot.send_photo(chat_id=user_id, photo=photo, caption=caption)
    except Exception as e:
        print(f"Failed to send message to {user_id}: {e}")


async def send_file_to_user(user_id: int, file: str, bot):
    try:
        document = FSInputFile(file)
        await bot.send_document(chat_id=user_id, document=document)
    except Exception as e:
        print(f"Failed to send message to {user_id}: {e}")


def create_language_keyboard(selected_languages):
    builder = ReplyKeyboardBuilder()

    for lang in button_languages:
        emoji = "❌" if lang in selected_languages else "✔️"
        builder.button(text=f"{emoji} {lang}")

    if selected_languages:
        builder.button(text="✅ Tasdiqlash")
        builder.button(text="Orqaga ↩️")
        builder.button(text="Bekor qilish 🚫")
        builder.adjust(3, 1, 2)
    else:
        builder.button(text="Orqaga ↩️")
        builder.button(text="Bekor qilish 🚫")
        builder.adjust(3, 2)
    return builder.as_markup(resize_keyboard=True)


def create_programms_keyboard(selected_programms):
    builder = ReplyKeyboardBuilder()

    for lang in button_programms:
        emoji = "❌" if lang in selected_programms else "✔️"
        builder.button(text=f"{emoji} {lang}")

    if selected_programms:
        builder.button(text="✅ Tasdiqlash")
        builder.button(text="Orqaga ↩️")
        builder.button(text="Bekor qilish 🚫")
        builder.adjust(3, 1, 2)
    else:
        builder.button(text="Orqaga ↩️")
        builder.button(text="Bekor qilish 🚫")
        builder.adjust(3, 2)
    return builder.as_markup(resize_keyboard=True)


def is_valid_phone_number(text):
    pattern = r"^\+998\d{9}$"
    return bool(re.match(pattern, text))


@router.message(CommandStart())
async def cm_start(message: Message):
    await message.answer(answer, reply_markup=kb.resume)


def build_keyboard(button_list):
    builder = ReplyKeyboardBuilder()

    for i in button_list:
        builder.button(text=i)

    builder.button(text="Orqaga ↩️")
    builder.button(text='Bekor qilish 🚫')
    if len(button_list) % 2 != 0:
        l = []
        for i in range((len(button_list) + 2) // 2):
            l.append(2)
        l[-1] = 1
        l.append(2)
        builder.adjust(*l)
    else:
        builder.adjust(2)
    return builder


def is_valid_date(date_text: str) -> bool:
    if not re.match(r"^\d{2}\.\d{2}\.\d{4}$", date_text):
        return False
    try:
        datetime.strptime(date_text, "%d.%m.%Y")
        return True
    except ValueError:
        return False


class Resume(StatesGroup):
    first_name = State()
    last_name = State()
    middle_name = State()
    gender = State()
    birthday = State()
    education = State()
    education_place = State()
    family_status = State()
    specialty = State()
    region = State()
    district = State()
    address = State()
    branch_desire = State()
    post = State()
    experience = State()
    worked_before = State()
    phone = State()
    extra_phone = State()
    shirt_size = State()
    is_studying = State()
    education_type = State()
    languages = State()
    selected_languages = State()
    work_experience = State()
    programms_experience = State()
    selected_programms = State()
    about = State()
    salary = State()
    is_familiar_works_here = State()
    is_uzbek_citizen = State()
    is_working_now = State()
    from_vacancy_info = State()
    face_photo = State()
    finish = State()


@router.message(F.text == "Anketa to'ldirish 📝")
async def first_name(message: Message, state: FSMContext):
    await state.set_state(Resume.first_name)
    await message.answer("👤 Ismingizni yozing:", reply_markup=kb.default)


@router.message(Resume.first_name)
async def last_name(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await message.answer("Siz birinchi bosqichdasiz.", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        await state.update_data(first_name=message.text)
        await state.set_state(Resume.last_name)
        await message.answer("👤 Familiyangizni yozing:", reply_markup=kb.default)


@router.message(Resume.last_name)
async def middle_name(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.first_name)
        await message.answer("👤 Ismingizni yozing:", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        await state.update_data(last_name=message.text)
        await state.set_state(Resume.middle_name)
        await message.answer("👤 Otangizni ismini yozing:", reply_markup=kb.default_with_skip)


@router.message(Resume.middle_name)
async def gender(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.last_name)
        await message.answer("👤 Familiyangizni yozing:", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == "▶️ O‘tkazib yuborish":
        await state.update_data(middle_name="-")
        await state.set_state(Resume.gender)
        await message.answer("🧑👩 Jinsingizni tanlang:", reply_markup=kb.gender)
    else:
        await state.update_data(middle_name=message.text)
        await state.set_state(Resume.gender)
        await message.answer("🧑👩 Jinsingizni tanlang:", reply_markup=kb.gender)


@router.message(Resume.gender)
async def birthday(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.middle_name)
        await message.answer("👤 Otangizni ismini yozing:", reply_markup=kb.default_with_skip)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ['Ayol', 'Erkak']:
        await state.update_data(gender=message.text)
        await state.set_state(Resume.birthday)
        await message.answer("📅 Tug‘ilgan sanangiz :\n\nKK.OO.YYYY(23.04.1998) formatida:", reply_markup=kb.default)
    else:
        await message.answer("🧑👩 Jinsingizni tanlang:", reply_markup=kb.gender)


@router.message(Resume.birthday)
async def education(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.gender)
        await message.answer("🧑👩 Jinsingizni tanlang:", reply_markup=kb.gender)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif is_valid_date(message.text):
        await state.update_data(birthday=message.text)
        await state.set_state(Resume.education)
        await message.answer("💼 Maʼlumotingizni tanlang:", reply_markup=kb.education)
    else:
        await message.answer("📅 Tug‘ilgan sanangiz :\n\nKK.OO.YYYY(23.04.1998) formatida:", reply_markup=kb.default)


@router.message(Resume.education)
async def education_place(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.birthday)
        await message.answer("📅 Tug‘ilgan sanangiz :\n\nKK.OO.YYYY(23.04.1998) formatida:", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ['Oliy', 'Magistratura', 'Talaba', "O'rta maxsus", "O'rta"]:
        await state.update_data(education=message.text)
        await state.set_state(Resume.education_place)
        await message.answer("Ta‘lim muassasasining nomi va bitirgan yilingiz:", reply_markup=kb.default_with_skip)
    else:
        await message.answer("💼 Maʼlumotingizni tanlang:", reply_markup=kb.education)


@router.message(Resume.education_place)
async def family_status(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.education)
        await message.answer("💼 Maʼlumotingizni tanlang:", reply_markup=kb.education)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == "▶️ O‘tkazib yuborish":
        await state.update_data(education_place="-")
        await state.set_state(Resume.family_status)
        await message.answer("👨‍👩‍👧‍👦 Oilaviy ahvolingiz:", reply_markup=kb.family_status)
    else:
        await state.update_data(education_place=message.text)
        await state.set_state(Resume.family_status)
        await message.answer("👨‍👩‍👧‍👦 Oilaviy ahvolingiz:", reply_markup=kb.family_status)


@router.message(Resume.family_status)
async def specialty(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.education_place)
        await message.answer("Ta‘lim muassasasining nomi va bitirgan yilingiz:", reply_markup=kb.default_with_skip)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ['Turmush qurgan', 'Turmush qurmagan']:
        await state.update_data(family_status=message.text)
        await state.set_state(Resume.specialty)
        await message.answer("👨‍🔧 Mutaxassisligingiz:", reply_markup=kb.default)
    else:
        await message.answer("👨‍👩‍👧‍👦 Oilaviy ahvolingiz:", reply_markup=kb.family_status)


@router.message(Resume.specialty)
async def region(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.family_status)
        await message.answer("👨‍👩‍👧‍👦 Oilaviy ahvolingiz:", reply_markup=kb.family_status)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        await state.update_data(specialty=message.text)
        await state.set_state(Resume.region)
        await message.answer("🌐 Yashash manzilingiz viloyat(xaqiqiy turar joy):", reply_markup=kb.region)


@router.message(Resume.region)
async def district(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.family_status)
        await message.answer("👨‍👩‍👧‍👦 Oilaviy ahvolingiz:", reply_markup=kb.family_status)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in list(regions.keys()):
        await state.update_data(region=message.text)
        await state.set_state(Resume.district)
        await message.answer("🌐 Yashash manzilingiz tuman(xaqiqiy turar joy):",
                             reply_markup=build_keyboard(regions[message.text]).as_markup(resize_keyboard=True,
                                                                                          one_time_keyboard=True))
    else:
        await message.answer("🌐 Yashash manzilingiz viloyat(xaqiqiy turar joy):", reply_markup=kb.region)


@router.message(Resume.district)
async def address(message: Message, state: FSMContext):
    data = await state.get_data()
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.region)
        await message.answer("🌐 Yashash manzilingiz viloyat(xaqiqiy turar joy):", reply_markup=kb.region)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in regions[data.get("region")]:
        await state.update_data(district=message.text)
        await state.set_state(Resume.address)
        await message.answer("🏘 To‘liq manzilingizni kiriting(MFY, ko‘cha):", reply_markup=kb.default)
    else:
        await message.answer("🌐 Yashash manzilingiz tuman(xaqiqiy turar joy):",
                             reply_markup=build_keyboard(regions[message.text]).as_markup(resize_keyboard=True,
                                                                                          one_time_keyboard=True))


@router.message(Resume.address)
async def branch_desire(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.district)
        await message.answer("🌐 Yashash manzilingiz tuman(xaqiqiy turar joy):",
                             reply_markup=build_keyboard(regions[message.text]).as_markup(resize_keyboard=True,
                                                                                          one_time_keyboard=True))
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        await state.update_data(address=message.text)
        await state.set_state(Resume.branch_desire)
        buttons = await get_branch()
        await message.answer("Siz qaysi filialda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))


@router.message(Resume.branch_desire)
async def post(message: Message, state: FSMContext):
    branches = await get_branch()
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.address)
        await message.answer("🏘 To‘liq manzilingizni kiriting(MFY, ko‘cha):", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in branches:
        await state.update_data(branch_desire=message.text)
        await state.set_state(Resume.post)
        buttons = await get_post()
        await message.answer("👨🏻‍💼 Qaysi lavozimda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))
    else:
        await message.answer("Siz qaysi filialda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(branches).as_markup(resize_keyboard=True,
                                                                             one_time_keyboard=True))


@router.message(Resume.post)
async def experience(message: Message, state: FSMContext):
    posts = await get_post()
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.branch_desire)
        buttons = await get_branch()
        await message.answer("Siz qaysi filialda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in posts:
        await state.update_data(post=message.text)
        await state.set_state(Resume.experience)
        await message.answer("Belgilangan lavozimdagi ish tajribangiz:", reply_markup=kb.experience)
    else:
        await message.answer("👨🏻‍💼 Qaysi lavozimda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(posts).as_markup(resize_keyboard=True,
                                                                          one_time_keyboard=True))


@router.message(Resume.experience)
async def worked_before(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.post)
        buttons = await get_post()
        await message.answer("👨🏻‍💼 Qaysi lavozimda ishlashni xohlaysiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ["Tajribam yo'q", '1 yildan 3 yilgacha', "3 yildan 6 yilgacha", '6 yildan yuqori']:
        await state.update_data(experience=message.text)
        await state.set_state(Resume.worked_before)
        await message.answer("Siz avval bizning kompaniyamizda ishlaganmisiz?", reply_markup=kb.yes_or_no)
    else:
        await message.answer("Belgilangan lavozimdagi ish tajribangiz:", reply_markup=kb.experience)


@router.message(Resume.worked_before)
async def phone(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.experience)
        await message.answer("Belgilangan lavozimdagi ish tajribangiz:", reply_markup=kb.experience)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ["Ha", 'Yo\'q']:
        await state.update_data(worked_before=message.text)
        await state.set_state(Resume.phone)
        await message.answer("📞 Telefon raqamingiz +998xxxxxxxxx formatida:", reply_markup=kb.phone)
    else:
        await message.answer("Siz avval bizning kompaniyamizda ishlaganmisiz?", reply_markup=kb.yes_or_no)


@router.message(Resume.phone)
async def extra_phone(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.worked_before)
        await message.answer("Siz avval bizning kompaniyamizda ishlaganmisiz?", reply_markup=kb.yes_or_no)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)

    elif message.contact is not None:
        await state.update_data(phone=message.contact.phone_number)
        await state.set_state(Resume.extra_phone)
        await message.answer("☎️ Qo‘shimcha telefon raqamingiz +998xxxxxxxxx formatida:",
                             reply_markup=kb.default_with_skip)
    elif is_valid_phone_number(message.text):
        await state.update_data(phone=message.text)
        await state.set_state(Resume.extra_phone)
        await message.answer("☎️ Qo‘shimcha telefon raqamingiz +998xxxxxxxxx formatida:",
                             reply_markup=kb.default_with_skip)
    else:
        await message.answer("📞 Telefon raqamingiz +998xxxxxxxxx formatida:", reply_markup=kb.phone)


@router.message(Resume.extra_phone)
async def is_studying(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.phone)
        await message.answer("📞 Telefon raqamingiz +998xxxxxxxxx formatida:", reply_markup=kb.phone)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == "▶️ O‘tkazib yuborish":
        await state.update_data(extra_phone="-")
        await state.set_state(Resume.is_studying)
        await message.answer("Siz hozirda qaysidir universitet, litsey yoki kollej talabasimisiz?",
                             reply_markup=kb.yes_or_no)

    elif is_valid_phone_number(message.text):
        await state.update_data(extra_phone=message.text)
        await state.set_state(Resume.is_studying)
        await message.answer("Siz hozirda qaysidir universitet, litsey yoki kollej talabasimisiz?",
                             reply_markup=kb.yes_or_no)
    else:
        await message.answer("☎️ Qo‘shimcha telefon raqamingiz +998xxxxxxxxx formatida:",
                             reply_markup=kb.default_with_skip)


@router.message(Resume.is_studying)
async def education_type(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.extra_phone)
        await message.answer("☎️ Qo‘shimcha telefon raqamingiz +998xxxxxxxxx formatida:",
                             reply_markup=kb.default_with_skip)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == 'Yo\'q':
        await state.update_data(is_studying=message.text)
        await state.update_data(education_type="-")
        await state.update_data(selected_languages=[])
        await state.set_state(Resume.languages)
        await message.answer("🇷🇺 Hozirda qaysi tillarni bilasiz?", reply_markup=create_language_keyboard([]))
    elif message.text == "Ha":
        await state.update_data(is_studying=message.text)
        await state.set_state(Resume.education_type)
        await message.answer("Qanday ta'lim shakli?", reply_markup=kb.education_type)
    else:
        await message.answer("Siz hozirda qaysidir universitet, litsey yoki kollej talabasimisiz?",
                             reply_markup=kb.yes_or_no)


@router.message(Resume.education_type)
async def languages(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.is_studying)
        await message.answer("Siz hozirda qaysidir universitet, litsey yoki kollej talabasimisiz?",
                             reply_markup=kb.yes_or_no)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ["Kunduzgi", 'Sirtqi']:
        await state.update_data(education_type=message.text)
        await state.update_data(selected_languages=[])
        await state.set_state(Resume.languages)
        await message.answer("🇷🇺 Hozirda qaysi tillarni bilasiz?", reply_markup=create_language_keyboard([]))
    else:
        await message.answer("Qanday ta'lim shakli?", reply_markup=kb.education_type)


@router.message(Resume.languages)
async def work_experience(message: Message, state: FSMContext):
    user_data = await state.get_data()
    selected_languages = user_data.get("selected_languages", [])
    language = message.text[2:].strip()
    if language in button_languages:
        if language in selected_languages:
            selected_languages.remove(language)
        else:
            selected_languages.append(language)

        await state.update_data(selected_languages=selected_languages)
        if any(selected_languages):
            s = ""
            n = 0
            for i in selected_languages:
                n += 1
                if n == len(selected_languages):
                    s += i
                else:
                    s += i + ", "
            await message.answer(
                "🇷🇺 Hozirda qaysi tillarni bilasiz?\n\n: " + s,
                reply_markup=create_language_keyboard(selected_languages)
            )
        else:
            await message.answer(
                "🇷🇺 Hozirda qaysi tillarni bilasiz?",
                reply_markup=create_language_keyboard(selected_languages)
            )
    elif message.text == "✅ Tasdiqlash" and selected_languages:
        await state.update_data(languages=selected_languages)
        await state.set_state(Resume.work_experience)
        await message.answer(
            "◀️🏦 Qayerda, qachon va kim bo'lib ishlaganingizni ayting. 3-4 o'rin uchun ahamiyatli bo'lgan ishlarni tavsiflang. Sizning rasmiy va norasmiy ish tajribangiz biz uchun muhim.Misol uchun, 2020-2022 yillarda 1Zumda do'konlar tarmog'ida rekruter",
            reply_markup=kb.default)
    elif message.text == "Orqaga ↩️":
        if user_data.get("is_studying") == 'Yo\'q':
            await state.set_state(Resume.is_studying)
            await message.answer("Siz hozirda qaysidir universitet, litsey yoki kollej talabasimisiz?",
                                 reply_markup=kb.yes_or_no)
        else:
            await state.set_state(Resume.education_type)
            await message.answer("Qanday ta'lim shakli?", reply_markup=kb.education_type)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        if any(selected_languages):
            s = ""
            n = 0
            for i in selected_languages:
                n += 1
                if n == len(selected_languages):
                    s += i
                else:
                    s += i + ", "
            await message.answer(
                "🇷🇺 Hozirda qaysi tillarni bilasiz?\n\n: " + s,
                reply_markup=create_language_keyboard(selected_languages)
            )
        else:
            await message.answer(
                "🇷🇺 Hozirda qaysi tillarni bilasiz?",
                reply_markup=create_language_keyboard(selected_languages)
            )


@router.message(Resume.work_experience)
async def programms_experience(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        user_data = await state.get_data()
        selected_languages = user_data.get("selected_languages", [])
        await state.set_state(Resume.languages)
        await message.answer("🇷🇺 Hozirda qaysi tillarni bilasiz?",
                             reply_markup=create_language_keyboard(selected_languages))
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)

    else:
        await state.update_data(work_experience=message.text)
        await state.set_state(Resume.programms_experience)
        await message.answer("👨‍💻 Qaysi dasturlardan foydalana olasiz?", reply_markup=create_programms_keyboard([]))


@router.message(Resume.programms_experience)
async def about(message: Message, state: FSMContext):
    user_data = await state.get_data()
    selected_programms = user_data.get("selected_programms", [])
    programm = message.text[2:].strip()
    if programm in button_programms:
        if programm in selected_programms:
            selected_programms.remove(programm)
        else:
            selected_programms.append(programm)

        await state.update_data(selected_programms=selected_programms)
        if any(selected_programms):
            s = ""
            n = 0
            for i in selected_programms:
                n += 1
                if n == len(selected_programms):
                    s += i
                else:
                    s += i + ", "
            await message.answer(
                "👨‍💻 Qaysi dasturlardan foydalana olasiz?\n\n: " + s,
                reply_markup=create_programms_keyboard(selected_programms)
            )
        else:
            await message.answer(
                "👨‍💻 Qaysi dasturlardan foydalana olasiz?",
                reply_markup=create_programms_keyboard(selected_programms)
            )
    elif message.text == "✅ Tasdiqlash" and selected_programms:
        await state.update_data(programms_experience=selected_programms)
        await state.set_state(Resume.about)
        await message.answer("📝 Qo‘shimcha maʼlumotlar:", reply_markup=kb.default_with_skip)

    elif message.text == "Orqaga ↩️":
        await state.set_state(Resume.work_experience)
        await message.answer(
            "◀️🏦 Qayerda, qachon va kim bo'lib ishlaganingizni ayting. 3-4 o'rin uchun ahamiyatli bo'lgan ishlarni tavsiflang. Sizning rasmiy va norasmiy ish tajribangiz biz uchun muhim.Misol uchun, 2020-2022 yillarda 1Zumda do'konlar tarmog'ida rekruter",
            reply_markup=kb.default)

    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)

    else:
        if any(selected_programms):
            s = ""
            n = 0
            for i in selected_programms:
                n += 1
                if n == len(selected_programms):
                    s += i
                else:
                    s += i + ", "
            await message.answer(
                "👨‍💻 Qaysi dasturlardan foydalana olasiz?\n\n: " + s,
                reply_markup=create_programms_keyboard(selected_programms)
            )
        else:
            await message.answer(
                "👨‍💻 Qaysi dasturlardan foydalana olasiz?",
                reply_markup=create_programms_keyboard(selected_programms)
            )


@router.message(Resume.about)
async def salary(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        user_data = await state.get_data()
        selected_programms = user_data.get("selected_programms", [])
        await state.set_state(Resume.programms_experience)
        await message.answer("👨‍💻 Qaysi dasturlardan foydalana olasiz?",
                             reply_markup=create_programms_keyboard(selected_programms))

    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == "▶️ O‘tkazib yuborish":
        await state.update_data(about="-")
        await state.set_state(Resume.salary)
        await message.answer("💰 Qancha maosh olishni xohlaysiz?(faqat raqamlarda kiriting):", reply_markup=kb.default)
    else:
        await state.update_data(about=message.text)
        await state.set_state(Resume.salary)
        await message.answer("💰 Qancha maosh olishni xohlaysiz?(faqat raqamlarda kiriting):", reply_markup=kb.default)


@router.message(Resume.salary)
async def is_familiar_works_here(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.about)
        await message.answer("📝 Qo‘shimcha maʼlumotlar:", reply_markup=kb.default_with_skip)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text.isdigit():
        await state.update_data(salary=message.text)
        await state.set_state(Resume.is_familiar_works_here)
        await message.answer(
            '❗️"1Zumda" kompaniyasida ishlaydigan yaqin qarindoshlaringiz bormi? Agar bo\'lsa, to\'liq familiyasi, ismi, otasining ismini va lavozimini yozing:',
            reply_markup=kb.yes_or_no)
    else:
        await message.answer("💰 Qancha maosh olishni xohlaysiz?(faqat raqamlarda kiriting):", reply_markup=kb.default)


@router.message(Resume.is_familiar_works_here)
async def is_uzbek_citizen(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.salary)
        await message.answer("💰 Qancha maosh olishni xohlaysiz?(faqat raqamlarda kiriting):", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    else:
        await state.update_data(is_familiar_works_here=message.text)
        await state.set_state(Resume.is_uzbek_citizen)
        await message.answer("Siz  O'zbekiston Respublikasi fuqarosimisiz?", reply_markup=kb.yes_or_no)


@router.message(Resume.is_uzbek_citizen)
async def is_working_now(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.is_familiar_works_here)
        await message.answer(
            '❗️"1Zumda" kompaniyasida ishlaydigan yaqin qarindoshlaringiz bormi? Agar bo\'lsa, to\'liq familiyasi, ismi, otasining ismini va lavozimini yozing:',
            reply_markup=kb.yes_or_no)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ["Ha", 'Yo\'q']:
        await state.update_data(is_uzbek_citizen=message.text)
        await state.set_state(Resume.is_working_now)
        await message.answer("Siz hozirda ish bilan ta'minlanganmisiz?", reply_markup=kb.yes_or_no)
    else:
        await message.answer("Siz  O'zbekiston Respublikasi fuqarosimisiz?", reply_markup=kb.yes_or_no)


@router.message(Resume.is_working_now)
async def from_vacancy_info(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.is_uzbek_citizen)
        await message.answer("Siz  O'zbekiston Respublikasi fuqarosimisiz?", reply_markup=kb.yes_or_no)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in ["Ha", 'Yo\'q']:
        buttons = await get_vacancy_info()

        await state.update_data(is_working_now=message.text)
        await state.set_state(Resume.from_vacancy_info)
        await message.answer("Bo'sh ish o'rni haqida qayerdan bildingiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))
    else:
        await message.answer("Siz  O'zbekiston Respublikasi fuqarosimisiz?", reply_markup=kb.yes_or_no)


@router.message(Resume.from_vacancy_info)
async def face_photo(message: Message, state: FSMContext):
    buttons = await get_vacancy_info()
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.is_working_now)
        await message.answer("Siz hozirda ish bilan ta'minlanganmisiz?", reply_markup=kb.yes_or_no)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text in buttons:
        await state.update_data(from_vacancy_info=message.text)
        await state.set_state(Resume.face_photo)
        await message.answer("📷 O‘zingizni rasmingizni yuboring:", reply_markup=kb.default)
    else:
        await message.answer("Bo'sh ish o'rni haqida qayerdan bildingiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))


@router.message(Resume.face_photo)
async def save_photo(message: Message, state: FSMContext):
    buttons = await get_vacancy_info()
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.from_vacancy_info)
        await message.answer("Bo'sh ish o'rni haqida qayerdan bildingiz?",
                             reply_markup=build_keyboard(buttons).as_markup(resize_keyboard=True,
                                                                            one_time_keyboard=True))
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.photo is not None:
        photo = message.photo[-1].file_id
        await state.update_data(face_photo=photo)
        data = await state.get_data()
        await message.answer_photo(data["face_photo"],
                                   caption=f'👤: {data["last_name"]} {data["first_name"]} {data["middle_name"]} ({data["gender"]})\n📆: {data["birthday"]}\n📍: {data["region"]}, {data["district"]}, {data["address"]}\n👨‍👩‍👧‍👦: {data["family_status"]}\n💼: {data["specialty"]}\n📞: {data["phone"]}\n🧳: {data["work_experience"]}\n🎓: {data["education"]}\n🏫: {data["education_place"]}\n🧑: {", ".join(data["programms_experience"])}\n🇷🇺🇺🇿🇺🇸: {", ".join(data["languages"])}\n🔍📍: {data["branch_desire"]}\n🧰: {data["post"]}\n💰: {data["salary"]}')
        await state.set_state(Resume.finish)
        await message.answer(
            "\"Men roziman\" tugmachasini bosish orqali siz o'zingizning shaxsiy ma'lumotlaringizni kompaniya maqsadlarida qayta ishlash uchun ularni saqlashga, foydalanishga va o'zaro almashishga rozilik bildirasiz. \nShuningdek, ushbu anketada siz taqdim etgan barcha ma'lumotlar ishonchli ekanligi va yolg'on ma'lumot uzatilishi holatlari aniqlangan taqdirda barcha javobgarlikni o'z zimmangizga olasiz.",
            reply_markup=kb.finish)


@router.message(Resume.finish)
async def finish(message: Message, state: FSMContext):
    if message.text == "Orqaga ↩️":
        await state.set_state(Resume.face_photo)
        await message.answer("📷 O‘zingizni rasmingizni yuboring:", reply_markup=kb.default)
    elif message.text == "Bekor qilish 🚫":
        await state.clear()
        await message.answer(
            "Salom 👋🏻\nUshbu bot 1Zumda labor️ anketalarni to‘ldirish va mehnat uchun mo‘ljallangan!\nBu yerda siz o‘zingizning arizangizni 📄 to‘ldirishingiz ✍️ va bizning kompaniyamizdagi mavjud bo‘sh ish o‘rinlari haqida bilib olishingiz mumkin!\n\nЗдравствуйте 👋🏻\nЭтот бот предназначен для заполнения анкеты ✍️ и трудоустройства  в 1Zumda!\nЗдесь Вы можете заполнить свою анкету 📄 и узнать о существующих вакансиях нашей Компании!",
            reply_markup=kb.resume)
    elif message.text == "Men roziman":
        data = await state.get_data()
        os.makedirs('djangobot/media', exist_ok=True)
        await message.answer(
            f"Hurmatli {data['last_name']} {data['first_name']}, vakansiyamizga va anketa so'rovnomasiga qiziqish bildirganingiz uchun tashakkur! Kompaniyamiz bilan hamkorlikni tiklashga tayyor ekanligingizdan xursandmiz. Sizning anketangizni xodimlarni ishga qayta tiklash bilan shug'ullanadigan bo'limga yuboramiz. Biz Sizga 3 ish kuni ichida ishga joylashish haqidagi arizangizga javob qaytaramiz. \n\nMUHIM! Iltimos, telegramingizni, ushbu botni va anketangizda ko'rsatilgan telefon raqamingizni o'chirmasligingizni so'raymiz. Biz Sizga ushbu bot yoki telefon orqali javob beramiz.\n\nHurmat bilan 1Zumda",
            reply_markup=ReplyKeyboardRemove())
        photo_id = data.get('face_photo')
        photo = await message.bot.get_file(photo_id)
        file_path = os.path.join('djangobot/media', data['last_name'] + " " + data['first_name'] + ".jpg")
        await photo.bot.download_file(photo.file_path, file_path)

        data['face_photo'] = file_path
        chat_ids = await get_chat_ids()
        for chat_id in chat_ids:
            await send_message_to_user(chat_id, photo_id,
                                       f'👤: {data["last_name"]} {data["first_name"]} {data["middle_name"]} ({data["gender"]})\n📆: {data["birthday"]}\n📍: {data["region"]}, {data["district"]}, {data["address"]}\n👨‍👩‍👧‍👦: {data["family_status"]}\n💼: {data["specialty"]}\n📞: {data["phone"]}\n🧳: {data["work_experience"]}\n🎓: {data["education"]}\n🏫: {data["education_place"]}\n🧑: {", ".join(data["programms_experience"])}\n🇷🇺🇺🇿🇺🇸: {", ".join(data["languages"])}\n🔍📍: {data["branch_desire"]}\n🧰: {data["post"]}\n💰: {data["salary"]}',
                                       message.bot)
            await send_file_to_user(chat_id, generate_resume(data), message.bot)
        await save_resume_data(data)
        await state.clear()
    else:
        await message.answer(
            "\"Men roziman\" tugmachasini bosish orqali siz o'zingizning shaxsiy ma'lumotlaringizni kompaniya maqsadlarida qayta ishlash uchun ularni saqlashga, foydalanishga va o'zaro almashishga rozilik bildirasiz. \nShuningdek, ushbu anketada siz taqdim etgan barcha ma'lumotlar ishonchli ekanligi va yolg'on ma'lumot uzatilishi holatlari aniqlangan taqdirda barcha javobgarlikni o'z zimmangizga olasiz.",
            reply_markup=kb.finish)
